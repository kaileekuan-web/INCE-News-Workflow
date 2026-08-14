#!/usr/bin/env python3
"""
Generate the AI News Word document.

The news table covers SMALLER AI STARTUPS. Stories about frontier labs and big
tech (OpenAI, Anthropic, Google/DeepMind, Meta, Microsoft, Amazon, Apple, xAI,
NVIDIA …) are dropped, as are opinion pieces, duplicates, and any link back to
x.com — see tools/news_filters.py for all four rules and `--include-frontier` /
`--include-opinion` to relax the first two.

The table renders flat, oldest first. The old OpenAI | Anthropic | BigTech |
Other grouping only makes sense when frontier coverage is present, so it comes
back with --include-frontier and is otherwise off — its classifier matches a
keyword anywhere in the text, which would bucket a Cohere post under "OpenAI"
for mentioning GPT.

Output: one news table, followed by a fundraising table.
Usage:
  python tools/generate_ai_doc.py --start_date 2026-03-31 --end_date 2026-04-09 \\
    --articles .tmp/summarized_articles.json [--chinese-only | --translate] \\
    --output-prefix AI_News
"""

import os
import sys
import json
import re
import argparse
import time
from io import BytesIO
from dotenv import load_dotenv

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tools.utils import format_date_for_display, call_claude, translate_to_chinese_claude
from tools.news_filters import curate, news_link
from tools.generate_word_doc import (
    add_hyperlink,
    add_formatted_text,
    extract_funding_with_web_search,
    create_funding_table,
    convert_bullets_to_paragraph,
    set_run_font,
    _set_para_font,
    pretranslate,
    translate_cached,
    FONT_ENGLISH,
    FONT_CHINESE,
)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import requests
except ImportError:
    print("ERROR: requests not installed. Run: pip install requests")
    sys.exit(1)


# ── Fundraising detection ──────────────────────────────────────────────────────


#tighten summarization
FUNDRAISING_KEYWORDS = [
    "raises", "raised", "raise", "funding", "fundrais",
    "series a", "series b", "series c", "series d", "series e",
    "seed round", "pre-seed", "venture", "investment round",
    "valuation", "unicorn", "acqui", "ipo", "goes public",
    "million", "billion", "$", "融资", "投资", "轮", "估值", "收购",
]


def is_fundraising_article(_article: dict) -> bool:
    """Disabled — all articles go to the main news table."""
    return False


def article_to_funding_event(article: dict) -> dict:
    """Use Claude to extract structured funding info from a detected funding article."""
    title = article.get("title", "")
    # Prefer full article content over the condensed summary so founder details aren't lost
    body = article.get("content") or article.get("summary") or article.get("description", "")
    date = article.get("published_at", "")[:10]

    prompt = (
        f"从以下新闻文章中提取融资信息，所有字段均用中文填写。\n\n"
        f"标题：{title}\n文章正文：{body}\n\n"
        f"返回一个JSON对象，包含以下字段：\n"
        f'"company": 获得融资的公司名称\n'
        f'"summary": 用中文描述该公司，包含：(1) 一句话说明公司核心业务，(2) 如文章中提到创始人背景信息，请附上（姓名、曾任职的公司和职责）。'
        f'参考格式："AI-native 网络安全公司，用 AI agent 实时检测攻击并自动响应。创始人 XX 曾负责 Amazon Web Services GuardDuty"\n'
        f'"stage": 融资轮次，用英文填写（Seed、Series A、Series B、Acquisition 等），未披露填 "Not disclosed"\n'
        f'"raise": 融资金额，用英文填写（例如 "US$50 million"），未披露填 "Not disclosed"\n'
        f'"valuation": 融资后估值，用英文填写（例如 "US$500 million"），未披露填 "Not disclosed"\n'
        f'"investors": 主要投资方，用英文填写（例如 "Sequoia Capital (lead), Index Ventures"），未披露填 "Not disclosed"\n\n'
        f"【事实约束】只能使用上方文章中明确出现的信息。不得根据公司名补全你已知的背景、"
        f"融资金额、估值、投资方或创始人经历；文章未写明的字段填 \"Not disclosed\"，不要猜测。\n\n"
        f"仅返回有效JSON，不要包含其他文字。"
    )

    text = call_claude(prompt, max_tokens=2048)
    if text:
        try:
            text = re.sub(r"^```json\s*", "", text)
            text = re.sub(r"\s*```$", "", text)
            data = json.loads(text)
            return {
                "date": date,
                "company": data.get("company", title),
                "summary": data.get("summary", ""),
                "stage": data.get("stage", "Not disclosed"),
                "raise": data.get("raise", "Not disclosed"),
                "valuation": data.get("valuation", "Not disclosed"),
                "investors": data.get("investors", "Not disclosed"),
                "_url": article.get("url", ""),
            }
        except Exception as e:
            print(f"    WARNING: extraction failed ({e}), using fallback")

    # Fallback: extraction failed
    raw = body.split(".")[0] if body else title
    return {
        "date": date,
        "company": title,
        "summary": raw[:200],
        "stage": "Not disclosed",
        "raise": "Not disclosed",
        "valuation": "Not disclosed",
        "investors": "Not disclosed",
        "_url": article.get("url", ""),
    }


# ── Group definitions ──────────────────────────────────────────────────────────

GROUP_ORDER = ["OpenAI", "Anthropic", "BigTech", "Other"]

GROUP_DISPLAY = {
    "OpenAI": "OpenAI",
    "Anthropic": "Anthropic",
    "BigTech": "大型科技公司（谷歌 · 苹果 · 亚马逊 · Meta · 微软 · Netflix · xAI · 英伟达）",
    "Other": "其他",
}

# Header row fill colours (light blue family)
GROUP_COLORS = {
    "OpenAI":    "D6E4F0",
    "Anthropic": "D6E4F0",
    "BigTech":   "D6E4F0",
    "Other":     "D6E4F0",
}

OPENAI_KEYWORDS = [
    "openai", "chatgpt", "gpt-4", "gpt-3", "gpt4", "gpt3", " gpt ",
    "sora", " o1 ", " o1\n", " o3 ", " o3\n", " o4 ", " o4\n",
    "dall-e", "dall·e", "whisper", "altman", "sam altman",
]
ANTHROPIC_KEYWORDS = ["anthropic", "claude"]
BIGTECH_KEYWORDS = [
    # Google
    "google", "谷歌", "deepmind", "gemini", "bard", "waymo",
    # Apple
    "apple", "苹果",
    # Amazon
    "amazon", "亚马逊", " aws ",
    # Meta
    " meta ", "meta\n", "facebook", "脸书", "instagram", "whatsapp", "llama",
    # Microsoft
    "microsoft", "微软", "bing", " azure ", "copilot",
    # Netflix
    "netflix", "奈飞",
    # xAI / Elon
    " xai ", "x.ai", "grok", "elon musk", "马斯克",
    # NVIDIA
    "nvidia", "英伟达", "cuda",
]


def is_mostly_chinese(text: str, threshold: float = 0.2) -> bool:
    """
    True if `text` already reads as Chinese.

    When the pipeline runs with --language zh the summariser has already
    produced Chinese prose, and translating it again is a wasted LLM call per
    article — a third of the run's entire request budget, which matters on a
    free tier with a daily cap. Titles stay English (they come from X and
    TechCrunch), so they still get translated.
    """
    if not text:
        return False
    cjk = sum(1 for ch in text if '一' <= ch <= '鿿')
    letters = sum(1 for ch in text if ch.isalpha() or '一' <= ch <= '鿿')
    return letters > 0 and (cjk / letters) >= threshold


def classify_article(article: dict) -> str:
    """Classify article into OpenAI | Anthropic | BigTech | Other by keyword match."""
    text = " " + " ".join([
        article.get("title", ""),
        article.get("summary", ""),
        article.get("description", ""),
    ]).lower() + " "

    for kw in OPENAI_KEYWORDS:
        if kw in text:
            return "OpenAI"
    for kw in ANTHROPIC_KEYWORDS:
        if kw in text:
            return "Anthropic"
    for kw in BIGTECH_KEYWORDS:
        if kw in text:
            return "BigTech"
    return "Other"


# ── Table helpers ──────────────────────────────────────────────────────────────

def _add_group_header_row(table, label: str, fill_hex: str = "D6E4F0"):
    """Add a full-width merged header row with coloured background."""
    row = table.add_row()
    # Merge both cells
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]

    # Background shading
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)

    # Text
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(label)
    run.bold = True
    set_run_font(run, font_size=11)
    run.font.color.rgb = RGBColor(31, 73, 125)


def create_grouped_news_table(
    doc: Document,
    articles: list,
    chinese_only: bool = False,
    translate: bool = False,
    group: bool = False,
):
    """
    Build the AI News table.

    `group=True` buckets articles OpenAI → Anthropic → BigTech → Other, which
    is only meaningful when frontier-lab coverage is in the report
    (--include-frontier). With the default scope those buckets are actively
    wrong: classify_article() matches a keyword anywhere in the text, so a
    Cohere post that mentions GPT lands under "OpenAI" — the passing mentions
    the curation step deliberately kept are exactly what it trips on. So the
    default is one flat table, oldest first.
    """
    # Counter for text that was already Chinese and needed no translation call.
    # Boxed in a list so the nested loop below can mutate it.
    skipped_translations = [0]

    # Translate everything up front, concurrently. This loop makes two calls per
    # article (title and summary); done inline they were two sequential network
    # round-trips per row, which on a 40-article report is most of the document
    # generation time. is_mostly_chinese() text is skipped exactly as before —
    # it is already Chinese and re-translating it is a wasted call.
    if chinese_only or translate:
        pending = []
        for article in articles:
            title = article.get("title", "无标题")
            if not is_mostly_chinese(title):
                pending.append(title)
            summary = convert_bullets_to_paragraph(
                article.get("summary", article.get("description", "暂无摘要")))
            if not is_mostly_chinese(summary):
                pending.append(summary)
        pretranslate(pending)

    # Classify and bucket
    groups = {g: [] for g in GROUP_ORDER}
    if group:
        for article in articles:
            groups[classify_article(article)].append(article)
    else:
        groups["Other"] = list(articles)

    # Sort within each group by date
    for g in GROUP_ORDER:
        groups[g].sort(key=lambda x: x.get("published_at", ""))

    total = sum(len(v) for v in groups.values())

    heading = doc.add_heading("AI 新闻摘要", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_font(heading)
    doc.add_paragraph(f"共 {total} 篇\n")

    # Create 2-column table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.columns[0].width = Inches(1.0)
    table.columns[1].width = Inches(6.0)

    # Header row
    hdr = table.rows[0].cells
    hdr[0].text = "日期"
    hdr[1].text = "摘要"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                set_run_font(run, font_size=12)

    # A single "其他" header over the whole table is noise, so headers only
    # appear when there is more than one group to tell apart.
    show_group_headers = sum(1 for g in GROUP_ORDER if groups[g]) > 1
    article_counter = 0
    for group in GROUP_ORDER:
        group_articles = groups[group]
        if not group_articles:
            continue

        if show_group_headers:
            _add_group_header_row(table, GROUP_DISPLAY[group], GROUP_COLORS[group])

        for article in group_articles:
            article_counter += 1
            row_cells = table.add_row().cells

            # Col 0: date
            date_str = format_date_for_display(article.get("published_at", ""))
            date_run = row_cells[0].paragraphs[0].add_run(date_str)
            set_run_font(date_run, font_size=10)

            # Col 1: hyperlinked title + summary
            summary_cell = row_cells[1]
            summary_para = summary_cell.paragraphs[0]

            title = article.get("title", "无标题")
            # Translate title to Chinese when in chinese-only or translate mode
            if chinese_only or translate:
                if is_mostly_chinese(title):
                    skipped_translations[0] += 1
                else:
                    title = translate_cached(title)
            # Link to the news the post pointed at, never to the post. A story
            # collected from a tweet that linked nowhere gets a bold headline
            # and no link — an x.com link is not a source.
            url = news_link(article)
            if url:
                add_hyperlink(summary_para, url, title)
            else:
                run = summary_para.add_run(title)
                run.bold = True
                set_run_font(run, font_size=10)

            # `or`, not a .get default — a failed summarization leaves an empty
            # string behind, and that would render as a blank cell.
            raw_summary = (article.get("summary") or article.get("description")
                           or article.get("title") or "暂无摘要")
            summary_text = convert_bullets_to_paragraph(raw_summary)

            set_run_font(summary_para.add_run("\n\n"), font_size=10)

            if chinese_only or translate:
                # With --language zh the summariser already wrote Chinese prose.
                # Re-translating it is one wasted inference per article — about a
                # third of the run — and on a local model that is minutes, not
                # cents. Titles still translate: they arrive in English.
                if is_mostly_chinese(summary_text):
                    skipped_translations[0] += 1
                    add_formatted_text(summary_para, summary_text, font_size=10)
                else:
                    add_formatted_text(summary_para, translate_cached(summary_text),
                                       font_size=10)
            else:
                add_formatted_text(summary_para, summary_text, font_size=10)

            # Embed article thumbnail after the summary text
            image_url = article.get("image_url", "")
            if image_url:
                try:
                    img_resp = requests.get(
                        image_url, timeout=5,
                        headers={"User-Agent": "Mozilla/5.0"},
                    )
                    img_resp.raise_for_status()
                    set_run_font(summary_para.add_run("\n"), font_size=10)
                    img_run = summary_para.add_run()
                    img_run.add_picture(BytesIO(img_resp.content), width=Inches(2.0))
                except Exception:
                    pass

    if skipped_translations[0]:
        print(f"  Skipped {skipped_translations[0]} translation call(s) — text was already Chinese")

    return table


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_ai_doc(
    start_date: str,
    end_date: str,
    articles_file: str = ".tmp/summarized_articles.json",
    output_dir: str = "output",
    max_articles: int = None,
    translate: bool = False,
    chinese_only: bool = True,
    output_prefix: str = "AI_News",
    no_funding: bool = False,
    funding_wechat_file: str = None,
    include_frontier: bool = False,
    include_opinion: bool = False,
    require_source: bool = True,
):
    load_dotenv(override=True)

    print("Loading data...")
    if not os.path.exists(articles_file):
        print(f"ERROR: {articles_file} not found")
        sys.exit(1)

    with open(articles_file, "r", encoding="utf-8") as f:
        articles = json.load(f)

    print(f"Loaded {len(articles)} articles")

    # Curate before --max, so a cap of 20 means 20 usable articles rather than
    # 20 candidates that filtering then cuts to six.
    print("Curating (startups only, news only, no duplicates)...")
    articles = curate(articles, include_frontier, include_opinion,
                      require_source=require_source)

    if max_articles:
        articles = articles[:max_articles]

    # Separate fundraising articles from regular news (unless disabled)
    if no_funding:
        regular_articles = articles
        funding_articles = []
    else:
        regular_articles = [a for a in articles if not is_fundraising_article(a)]
        funding_articles = [a for a in articles if is_fundraising_article(a)]
        if funding_articles:
            print(f"  Detected {len(funding_articles)} fundraising articles — moving to funding table")

    if translate:
        print("Translation enabled (Claude)")
    elif chinese_only:
        print("Chinese mode enabled (Claude for title translation)")

    anthropic_key = os.getenv("ANTHROPIC_API_KEY")

    print("Creating Word document...")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    # Stamp both ASCII and East Asian fonts onto the Normal style's XML so CJK
    # characters don't fall back to Word's default (宋体/SimSun) in compatibility mode.
    _rPr = style.element.find(qn('w:rPr'))
    if _rPr is None:
        _rPr = OxmlElement('w:rPr')
        style.element.append(_rPr)
    _rFonts = _rPr.find(qn('w:rFonts'))
    if _rFonts is None:
        _rFonts = OxmlElement('w:rFonts')
        _rPr.insert(0, _rFonts)
    _rFonts.set(qn('w:ascii'), FONT_ENGLISH)
    _rFonts.set(qn('w:hAnsi'), FONT_ENGLISH)
    _rFonts.set(qn('w:eastAsia'), FONT_CHINESE)
    _rFonts.set(qn('w:cs'), FONT_ENGLISH)
    for _attr in (qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')):
        if _attr in _rFonts.attrib:
            del _rFonts.attrib[_attr]

    title_para = doc.add_heading("AI 资讯报告", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(title_para)

    subtitle = doc.add_paragraph(f"{start_date}  至  {end_date}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_run_font(run, font_size=14)
        run.font.color.rgb = RGBColor(128, 128, 128)

    total_str = str(len(regular_articles))
    if not no_funding and funding_articles:
        total_str += f"（另有 {len(funding_articles)} 篇移至融资动态）"
    doc.add_paragraph(f"文章总数：{total_str}")
    doc.add_paragraph("")

    print("Creating news table...")
    create_grouped_news_table(doc, regular_articles, chinese_only, translate,
                              group=include_frontier)

    if not no_funding:
        # Load WeChat fundraising articles if provided
        wechat_funding = []
        if funding_wechat_file and os.path.exists(funding_wechat_file):
            with open(funding_wechat_file, "r", encoding="utf-8") as f:
                wechat_funding_articles = json.load(f)
            print(f"Extracting funding details from {len(wechat_funding_articles)} WeChat fundraising article(s)...")
            for i, a in enumerate(wechat_funding_articles, 1):
                print(f"  [{i}/{len(wechat_funding_articles)}] {a.get('title', '')[:70]}...")
                wechat_funding.append(article_to_funding_event(a))

        print("Searching for AI funding news with Claude web search (day by day)...")
        if anthropic_key:
            funding_events = extract_funding_with_web_search(anthropic_key, start_date, end_date)
            all_funding = funding_events + wechat_funding
            print(f"  Found {len(funding_events)} from web search + {len(wechat_funding)} from WeChat")
            create_funding_table(doc, all_funding)
        elif wechat_funding:
            print("  ANTHROPIC_API_KEY not set — using WeChat fundraising articles only")
            create_funding_table(doc, wechat_funding)
        else:
            print("  WARNING: ANTHROPIC_API_KEY not set — skipping funding section")

    os.makedirs(output_dir, exist_ok=True)
    filename = f"{output_prefix}_{start_date.replace('-','')}_{end_date.replace('-','')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    print(f"\n✓ Document saved to {filepath}")
    return filepath


def main():
    parser = argparse.ArgumentParser(description="Generate grouped AI News Word document")
    parser.add_argument("--start_date", required=True, help="Start date (YYYY-MM-DD)")
    parser.add_argument("--end_date", required=True, help="End date (YYYY-MM-DD)")
    parser.add_argument("--articles", default=".tmp/summarized_articles.json")
    parser.add_argument("--output_dir", default="output")
    parser.add_argument("--max", type=int, default=None)
    parser.add_argument("--translate", action="store_true", help="Add Chinese translation")
    parser.add_argument("--chinese-only", action="store_true", help="Chinese summary only")
    parser.add_argument("--no-funding", action="store_true", help="Skip funding detection and funding table")
    parser.add_argument("--output-prefix", default="AI_News")
    parser.add_argument("--funding-wechat", default=None, help="Path to summarized WeChat fundraising articles JSON")
    parser.add_argument("--include-frontier", action="store_true",
                        help="Keep stories about frontier labs and big tech "
                             "(frontier_labs.txt). Off by default: the news table "
                             "is about smaller AI startups.")
    parser.add_argument("--include-opinion", action="store_true",
                        help="Keep commentary, opinion pieces and executive statements. "
                             "Off by default: the news table is for reported events.")
    parser.add_argument("--allow-unlinked", dest="require_source",
                        action="store_false",
                        help="Keep articles with no link to published coverage, as "
                             "unlinked headlines. By default they are dropped — run "
                             "tools/resolve_sources.py first to find their coverage.")
    args = parser.parse_args()

    generate_ai_doc(
        args.start_date,
        args.end_date,
        args.articles,
        args.output_dir,
        args.max,
        args.translate,
        args.chinese_only,
        args.output_prefix,
        args.no_funding,
        args.funding_wechat,
        args.include_frontier,
        args.include_opinion,
        args.require_source,
    )


if __name__ == "__main__":
    main()
