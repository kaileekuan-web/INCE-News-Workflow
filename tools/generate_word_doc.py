#!/usr/bin/env python3
"""
Generate Word document with AI News table:
- Title (hyperlinked) + Date + Source | Summary (with optional Chinese translation)

Editorial rules applied before anything is rendered (tools/news_filters.py):
duplicates collapsed always; for --funding-topic ai, stories about frontier
labs / big tech and opinion pieces are dropped, so the table is about smaller
AI startups. Headlines link to the news a post pointed at — never to x.com.
Relax with --include-frontier / --include-opinion.

Uses python-docx for formatting
"""

import os
import sys
import json
import re
import argparse
import time
from datetime import datetime
from urllib.parse import urlsplit
from dotenv import load_dotenv

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from tools.utils import (
    format_date_for_display,
    call_claude,
    call_claude_search,
    translate_to_chinese_claude,
    parallel_map,
    MAX_WORKERS,
)
from tools.detect_funding import extract_funding_with_claude
from tools.verify_emit import emit_funding_claims
from tools.news_filters import (
    curate, news_link, is_aggregator, canonical_url, clean_headline,
    filter_news_section, signal_score,
)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

FONT_ENGLISH = "Microsoft YaHei"
FONT_CHINESE = "微软雅黑"


def set_run_font(run, font_size: int = 10):
    """Set consistent ASCII and East Asian fonts on a run."""
    run.font.size = Pt(font_size)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    # Set explicit fonts
    rFonts.set(qn('w:ascii'), FONT_ENGLISH)
    rFonts.set(qn('w:hAnsi'), FONT_ENGLISH)
    rFonts.set(qn('w:eastAsia'), FONT_CHINESE)
    rFonts.set(qn('w:cs'), FONT_ENGLISH)
    # Remove theme font attributes — in OOXML, theme attributes take precedence
    # over explicit font names and will cause inconsistency if left in place.
    for attr in (qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')):
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]


def _set_para_font(para):
    """Set consistent font names on every run in a paragraph without touching font size.
    Used for headings and plain paragraphs so the style-defined size is preserved."""
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), FONT_ENGLISH)
        rFonts.set(qn('w:hAnsi'), FONT_ENGLISH)
        rFonts.set(qn('w:eastAsia'), FONT_CHINESE)
        rFonts.set(qn('w:cs'), FONT_ENGLISH)
        for attr in (qn('w:asciiTheme'), qn('w:hAnsiTheme'), qn('w:eastAsiaTheme'), qn('w:cstheme')):
            if attr in rFonts.attrib:
                del rFonts.attrib[attr]

def _set_table_cell_margins(table, top: int = 80, bottom: int = 80,
                             left: int = 100, right: int = 100):
    """Set uniform cell padding (twips) on all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblCellMar'))
    if existing is not None:
        tblPr.remove(existing)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)
    tblPr.append(tblCellMar)


def _add_horizontal_rule(doc: Document, color: str = '4472C4'):
    """Add a thin paragraph-border line as a visual section divider."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


try:
    import requests
except ImportError:
    print("ERROR: requests not installed. Run: pip install requests")
    sys.exit(1)


def add_hyperlink(paragraph, url: str, text: str, font_size: int = 10):
    """
    Add a hyperlink to a paragraph

    python-docx doesn't have native hyperlink support, so we use XML

    Args:
        paragraph: docx paragraph object
        url: URL to link to
        text: Link text
        font_size: Font size in points
    """
    # Create hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element
    run = OxmlElement('w:r')

    # Run properties (style)
    rPr = OxmlElement('w:rPr')

    # Font names (ASCII + East Asian)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_ENGLISH)
    rFonts.set(qn('w:hAnsi'), FONT_ENGLISH)
    rFonts.set(qn('w:eastAsia'), FONT_CHINESE)
    rFonts.set(qn('w:cs'), FONT_ENGLISH)
    rPr.append(rFonts)

    # Font size (half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(font_size * 2))
    rPr.append(szCs)

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    run.append(text_elem)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_formatted_text(paragraph, text: str, font_size: int = 10):
    """
    Add text to paragraph with markdown bold (**text**) converted to Word bold

    Args:
        paragraph: docx paragraph object
        text: Text that may contain **bold** markdown
        font_size: Font size in points
    """
    # Pattern to match **bold** text
    pattern = r'\*\*(.+?)\*\*'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Add text before the bold part
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            set_run_font(run, font_size=font_size)

        # Add bold text
        run = paragraph.add_run(match.group(1))
        run.bold = True
        set_run_font(run, font_size=font_size)

        last_end = match.end()

    # Add remaining text after last match
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        set_run_font(run, font_size=font_size)


# Shape of the funding search reply. Constrained decoding guarantees it, which
# replaces the old "regex a JSON array out of the prose" step — that step was the
# part most likely to break silently as model output style drifts.
FUNDING_EVENTS_SCHEMA = {
    "type": "object",
    "properties": {
        "events": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "date": {
                        "type": "string",
                        "description": "Date the round was FIRST publicly announced, YYYY-MM-DD",
                    },
                    "announced_time": {
                        "type": "string",
                        "description": "Time of the announcement if the source gives one "
                                       "(HH:MM, 24h, any timezone), else empty string. "
                                       "Used only to order same-day announcements.",
                    },
                    "company": {"type": "string", "description": "Company name, in English"},
                    "summary": {
                        "type": "string",
                        "description": "2-3 sentences of professional Simplified Chinese",
                    },
                    "stage": {
                        "type": "string",
                        "description": "English, e.g. Seed, Series A, Acquisition, "
                                       "or 'Not disclosed'",
                    },
                    "raise": {
                        "type": "string",
                        "description": "English, e.g. 'US$50 million', or 'Not disclosed'",
                    },
                    "valuation": {
                        "type": "string",
                        "description": "English post-money valuation, or 'Not disclosed'",
                    },
                    "investors": {
                        "type": "string",
                        "description": "English, lead first, e.g. "
                                       "'Sequoia Capital (lead), Index Ventures'",
                    },
                    "url": {
                        "type": "string",
                        "description": "Primary source: the original announcement, "
                                       "or the most authoritative report of it",
                    },
                    "sources": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "Every URL used to verify this round, primary first. "
                                       "Two or more whenever they exist.",
                    },
                    "discrepancy": {
                        "type": "string",
                        "description": "If sources conflict on amount, stage or investors, "
                                       "state the conflict in English and name the sources. "
                                       "Empty string when they agree.",
                    },
                },
                "required": ["date", "announced_time", "company", "summary", "stage",
                             "raise", "valuation", "investors", "url", "sources",
                             "discrepancy"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["events"],
    "additionalProperties": False,
}


# One web search per day found the rounds one feed happened to carry, and
# nothing else. The week of 2026-08-10 is the worked example: Echovane, Point2
# Technology, Discovered Materials, River AI, Legio, Palette, CodeRabbit, Skan
# AI, Yuno, Remepy, Silicon Data and Mindgard all announced inside the range and
# none of them reached the table. Every one of them was public somewhere — a
# database entry, a Form D, an investor's own post — just not in the one place
# the single search looked.
#
# So the day is searched several times over, each pass pointed at a different
# kind of source. They overlap heavily by design: a round that shows up in three
# passes is a round three independent places agree on, and _merge_funding_events
# collapses the copies into the entry with the most disclosed fields.
FUNDING_SOURCE_PASSES = [
    {
        'key': 'primary',
        'label': 'company & investor announcements',
        'instruction': (
            "Search the companies' OWN announcements and their investors' own posts: "
            "company blogs, /news and /press pages, press releases on Business Wire, "
            "PR Newswire, GlobeNewswire and EIN Presswire, and the announcement posts "
            "VC firms publish about rounds they led (a16z, Sequoia, Index, Accel, "
            "Benchmark, Lightspeed, General Catalyst, Khosla, Greylock, Bessemer, "
            "First Round, Founders Fund, Kleiner Perkins, Insight, Y Combinator, "
            "Antler, Seedcamp and their peers). These are the primary sources — "
            "prefer them over anyone's coverage of them."
        ),
    },
    {
        'key': 'newsletters',
        'label': 'venture deal newsletters',
        'instruction': (
            "Search the daily venture deal newsletters and their deal logs: Axios Pro "
            "Rata, Fortune Term Sheet, StrictlyVC, PE Hub, Dan Primack's feed, "
            "Newcomer, The Information's venture coverage. These carry the small "
            "rounds trade press does not write up individually — read the deal lists "
            "at the bottom of each issue, not only the lead story."
        ),
    },
    {
        'key': 'databases',
        'label': 'funding databases',
        'instruction': (
            "Search the funding databases and their news arms: Crunchbase News and "
            "Crunchbase's daily funding roundups, PitchBook news, Dealroom, Tracxn, "
            "CB Insights. Use their per-day funding listings, which are compiled "
            "rather than reported and so include rounds nobody wrote an article about."
        ),
    },
    {
        'key': 'filings',
        'label': 'regulatory filings',
        'instruction': (
            "Search regulatory filings for rounds disclosed by filing rather than by "
            "announcement: SEC Form D filings on EDGAR (Rule 506 offerings), and the "
            "equivalent registries elsewhere. A Form D filed on the date is a "
            "disclosed round even when no publication covered it — report it, with "
            "the filing as the source, and mark undisclosed fields \"Not disclosed\" "
            "rather than inferring them."
        ),
    },
    {
        'key': 'trade',
        'label': 'trade press',
        'instruction': (
            "Search the trade press funding desks: TechCrunch's funding and venture "
            "tags, VentureBeat, Sifted, Tech.eu, EU-Startups, FinSMEs, Business "
            "Insider, Reuters, Bloomberg, Fortune, and the regional outlets (Calcalist "
            "and Globes for Israel, Nikkei and TechNode for Asia, LatamList for Latin "
            "America). Include non-English coverage."
        ),
    },
]

FUNDING_PASS_KEYS = [p['key'] for p in FUNDING_SOURCE_PASSES]


def _search_funding_single_day(api_key: str, date: str, topic: str,
                               source_pass: dict = None) -> list:
    """
    Search for funding events on a single date using Claude's server-side web
    search tool. Returns list of funding event dicts.

    `source_pass` is one entry from FUNDING_SOURCE_PASSES — the kind of source
    this pass should look in. None searches broadly, which is what the self-test
    and any direct caller get.

    api_key is accepted for call-site compatibility and ignored — the Anthropic
    client reads ANTHROPIC_API_KEY from the environment.
    """
    if topic == 'deeptech':
        sector = ("deep tech startups (robotics, advanced materials, quantum computing, "
                  "bio/medtech, space, semiconductors, clean energy)")
        sector_rule = ("INCLUDE hard-tech companies in those fields. EXCLUDE pure software "
                       "plays with no deep-tech component.")
    else:
        sector = "AI startups"
        sector_rule = (
            "INCLUDE companies whose core technology is AI: LLMs, generative AI, AI agents, "
            "computer vision, voice AI, AI infrastructure, AI-native SaaS.\n"
            "EXCLUDE companies where AI is a peripheral feature, and exclude conventional "
            "data storage, non-AI cybersecurity, general cloud infrastructure, "
            "blockchain/crypto and conventional fintech."
        )

    if source_pass:
        where = (f"Where to search (this pass covers {source_pass['label']} only):\n"
                 f"{source_pass['instruction']}\n\n"
                 "Other passes cover the other kinds of source, so search this kind "
                 "thoroughly rather than broadly. Small rounds count as much as large "
                 "ones: a $3M seed nobody wrote an article about is exactly what this "
                 "pass exists to find. Do not skip a round for being minor, for being "
                 "outside the US, or for having only one source.")
    else:
        where = ("Search TechCrunch, Reuters, Bloomberg, Fortune, VentureBeat, "
                 "Crunchbase News, Sifted, Axios Pro Rata, The Information, SEC Form D "
                 "filings, and official company press releases/blogs.")

    prompt = f"""Task: Find {sector} fundraising announcements that were FIRST publicly announced on {date}.

{where}

Consolidate duplicate reports of the same funding round into a single entry.

Requirements:
- Use {date} as the filtering criterion. Include only funding rounds that were FIRST publicly announced on that date.
- Filter by the funding ANNOUNCEMENT date, NOT the publication date of recap articles or later news coverage.
- {sector_rule}
- Do NOT include:
  - Funding rounds announced before or after {date}.
  - Articles that recap previously announced funding rounds.
  - Unicorn lists, funding roundups, market overviews, or trend articles.
  - Rumors, fundraising discussions, or unconfirmed reports.
- If no qualifying fundraising announcements exist for {date}, return an empty events array. Do NOT substitute companies from nearby dates — a near-miss silently widens the report's date range, which is worse than an empty day.

Verification:
- Verify each round against at least two reputable sources whenever they exist, and list every URL you used in "sources", primary first.
- "Primary" means the announcement itself — the company's own blog post, press release or /news page, or the lead investor's own post about the round. Put that URL in "url" and first in "sources". Use a trade-press or database writeup as "url" only when no primary source exists. A round that a company announced on its own blog should not be sourced to a summary of that blog post.
- One credible source is enough to REPORT a round; two are for confirming its details. Do not drop a round because you found only the filing, only the database entry, or only the investor's post — report it with what that source states and "Not disclosed" for the rest.
- If sources conflict (amount, stage, investors), still report the round using the most authoritative source, and describe the conflict in "discrepancy", naming the sources.
- If a detail is not publicly disclosed, write exactly "Not disclosed". Never estimate, and never carry a figure over from a different round.
- "Not disclosed" means no source states it — not that you did not look. If a source names the round (Seed, Series A/B/C) or the lead investor, report it. Read the sources you already found before falling back to "Not disclosed" on stage or investors.
- Every field must be supported by a source you found in this search. Do not fill gaps from prior knowledge.

Language:
- company, stage, raise, valuation, investors: English. Amounts as "US$50 million".
- summary: professional Simplified Chinese ONLY, 2-3 sentences, suitable for a VC investment newsletter. Cover what the company does, the round itself, and — if disclosed — how the capital will be used.
- Factual, objective, information-dense. No marketing language.
- Vary how each summary opens. Do NOT begin with "这是一家…". Start instead from the product, the technology, the target market, the financing event, or the use of funds — for example "<Company> 专注于……", "<Company> 开发……", "本轮融资将用于……", "面向……市场，<Company>……". Each summary should differ in structure from the others.

Return one object per distinct funding round."""

    pass_key = source_pass['key'] if source_pass else 'broad'
    result = call_claude_search(prompt, schema=FUNDING_EVENTS_SCHEMA,
                                label=f"funding search {date} [{pass_key}]")
    if not result:
        return []
    events = result.get('events')
    if not isinstance(events, list):
        return []
    # Stamp the pass onto each event so the merge can report how many
    # independent kinds of source found the same round.
    for event in events:
        if isinstance(event, dict):
            event['_passes'] = [pass_key]
    return events


def _filter_ai_funding_events(events: list) -> list:
    """
    Use Claude to validate that each funding event is genuinely AI-focused.
    Batches all companies in one call. Returns the filtered list.
    """
    if not events:
        return events

    companies = [
        {"index": i, "company": e.get('company', ''), "summary": e.get('summary', '')}
        for i, e in enumerate(events)
    ]

    prompt = (
        "以下是一批融资新闻公司，请判断每家公司是否以AI/人工智能为核心业务。\n\n"
        "判断标准：\n"
        "- 是：公司的主要产品/服务以AI/ML为技术基础（大语言模型、生成式AI、AI agent、"
        "计算机视觉、AI-native SaaS等）\n"
        "- 否：AI只是次要功能，或公司属于传统数据存储、传统网络安全、"
        "普通云基础设施、区块链/加密、传统金融科技\n\n"
        f"公司列表：\n{json.dumps(companies, ensure_ascii=False)}\n\n"
        "返回一个JSON数组，只包含确实是AI公司的index值。例如：[0, 2, 4]\n"
        "只返回JSON数组，不含其他文字。"
    )

    text = call_claude(prompt, max_tokens=2048)
    if text:
        match = re.search(r'\[[\d,\s]*\]', text)
        if match:
            try:
                keep_indices = set(json.loads(match.group()))
                filtered = [e for i, e in enumerate(events) if i in keep_indices]
                removed = len(events) - len(filtered)
                if removed:
                    print(f"  Validation removed {removed} non-AI company/companies")
                return filtered
            except Exception as e:
                print(f"  WARNING: AI validation failed ({e}), keeping all events")

    return events


def _company_key(name: str) -> str:
    """Normalize a company name to a dedup key.

    Strips Chinese characters first so that bilingual names like
    "谷歌母公司 Alphabet" and "Alphabet" collapse to the same key.
    Falls back to the full lowercased name for Chinese-only companies.
    """
    english = re.sub(r'[一-鿿　-〿＀-￯\s]+', ' ', name).strip().lower()
    return english if english else name.lower().strip()


def _merge_funding_events(base: list, extra: list) -> list:
    """Merge two funding event lists, deduplicating by company name.

    'base' events take priority — they come from Claude's extraction of our
    collected articles (higher fidelity). 'extra' events from the Claude web
    search are only added if the company wasn't already found in base.
    When two entries for the same company exist, the one with more disclosed
    fields is kept. This is also what consolidates the same round reported by
    several publications into one row.
    """
    by_company: dict = {}
    for e in base + extra:
        raw_name = e.get('company', '')
        if not raw_name:
            continue
        rounds = by_company.setdefault(_company_key(raw_name), [])
        for i, existing in enumerate(rounds):
            if _same_round(existing, e):
                rounds[i] = _consolidate(existing, e)
                break
        else:
            rounds.append(e)
    return [r for rounds in by_company.values() for r in rounds]


def _round_amount(event: dict) -> str:
    """Digits of the raise amount, for comparing 'US$20 million' with '$20M'."""
    raw = str(event.get('raise') or '')
    if not _is_disclosed(raw):
        return ''
    match = re.search(r'(\d[\d,.]*)\s*(k|m|b|bn|thousand|million|billion|万|亿)?', raw, re.I)
    if not match:
        return ''
    unit = (match.group(2) or '').lower()
    return f"{match.group(1).replace(',', '')}{unit[:1]}"


def _same_round(a: dict, b: dict) -> bool:
    """
    Do two entries for the same company describe the SAME funding round?

    Company name alone is not enough. A seed round in March and a Series B in
    August are one company and two events, and merging them silently drops a
    round and files the survivor under the wrong date and stage.

    Agreement on stage or amount settles it. Failing both — which happens when
    one report leaves them undisclosed — entries announced within three days of
    each other are treated as one round, since that is a recap window, not a
    second raise.
    """
    stage_a, stage_b = str(a.get('stage') or ''), str(b.get('stage') or '')
    if _is_disclosed(stage_a) and _is_disclosed(stage_b):
        if stage_a.strip().lower() != stage_b.strip().lower():
            return False            # different stages: different rounds
    amount_a, amount_b = _round_amount(a), _round_amount(b)
    if amount_a and amount_b and amount_a != amount_b:
        return False                # different amounts: different rounds
    if (amount_a and amount_a == amount_b) or (
            _is_disclosed(stage_a) and stage_a.strip().lower() == stage_b.strip().lower()):
        return True

    date_a, date_b = (a.get('date') or '')[:10], (b.get('date') or '')[:10]
    if date_a and date_b:
        try:
            gap = abs((datetime.strptime(date_a, '%Y-%m-%d')
                       - datetime.strptime(date_b, '%Y-%m-%d')).days)
            return gap <= 3
        except ValueError:
            pass
    return True                     # too little to tell them apart; treat as one


# Fields that describe the round itself. On consolidation each one is taken from
# whichever copy actually discloses it.
_FUNDING_FACT_FIELDS = ('stage', 'raise', 'valuation', 'investors', 'summary',
                        'company', 'url', 'announced_time')


def _consolidate(a: dict, b: dict) -> dict:
    """
    Fold two reports of one funding round into a single entry.

    The richer copy is the base, then anything it leaves undisclosed is filled
    from the other — two outlets rarely disclose the same subset, and taking one
    entry wholesale threw away whichever details only the other one carried.

    The date is the EARLIEST of the two, always. This is the rule the whole
    brief turns on: a recap published two days later is not a new announcement,
    and letting the recap's date win because its article happened to name the
    valuation would file the round under the wrong day.
    """
    winner, loser = (b, a) if _disclosed_count(b) > _disclosed_count(a) else (a, b)
    merged = dict(winner)

    for field in _FUNDING_FACT_FIELDS:
        if not _is_disclosed(merged.get(field)) and _is_disclosed(loser.get(field)):
            merged[field] = loser[field]

    dates = [d for d in (a.get('date'), b.get('date')) if d]
    if dates:
        merged['date'] = min(dates)

    merged['sources'] = _merge_sources(winner, loser)
    if not (merged.get('discrepancy') or '').strip():
        merged['discrepancy'] = loser.get('discrepancy', '') or ''

    # Which source passes found this round, unioned — the corroboration count.
    passes = []
    for e in (a, b):
        for key in (e.get('_passes') or []):
            if key not in passes:
                passes.append(key)
    if passes:
        merged['_passes'] = passes
    return merged


# Values meaning "the source did not say". "Not disclosed" is what the funding
# prompt asks for; the Chinese spellings are kept so entries from an older run,
# or from the article-extraction path, still score correctly here.
UNDISCLOSED = {'not disclosed', '不详', '未披露', '未公开', 'n/a', 'na', 'none',
               'unknown', '', 'null'}


def _is_disclosed(value) -> bool:
    return str(value).strip().lower() not in UNDISCLOSED


def _disclosed_count(event: dict) -> int:
    return sum(1 for k, v in event.items()
               if k not in ('sources', 'discrepancy', '_passes') and _is_disclosed(v))


def _merge_sources(*events) -> list:
    """Union of every source URL across copies of one round, order preserved."""
    out = []
    for e in events:
        urls = e.get('sources') or []
        if isinstance(urls, str):
            urls = [urls]
        for u in list(urls) + [e.get('url', '')]:
            u = (u or '').strip()
            if u and u not in out:
                out.append(u)
    return out


def _prefer_primary_source(event: dict) -> dict:
    """
    Reorder a round's sources so the announcement itself comes first, and point
    `url` at it.

    The company's own post and the investor's own post are the round; TechCrunch
    and Crunchbase are reports of it. When both are in hand the table should link
    to the former — the numbers are first-hand there, it is rarely paywalled, and
    it is the thing every writeup is downstream of. When only coverage exists,
    nothing changes: an aggregator link is still a source.
    """
    sources = _merge_sources(event)
    if not sources:
        return event

    seen, ordered = set(), []
    for url in sources:
        key = canonical_url(url)
        if key in seen:
            continue
        seen.add(key)
        ordered.append(url)

    # Stable partition, not a sort by rank — within each group the model's own
    # ordering is the best signal available about which source it trusted.
    primary = [u for u in ordered if not is_aggregator(u)]
    secondary = [u for u in ordered if is_aggregator(u)]

    out = dict(event)
    out['sources'] = primary + secondary
    out['url'] = out['sources'][0]
    return out


def extract_funding_with_web_search(api_key: str, start_date: str, end_date: str,
                                    topic: str = 'ai', passes: list = None) -> list:
    """
    Use Claude with server-side web search to find funding events in a date range.

    Searches day-by-day — the model skips dates in long ranges — and each day
    several times over, once per entry in FUNDING_SOURCE_PASSES. See the comment
    there for why: one pass finds the rounds one kind of source happened to
    carry, and the rounds it misses are not random. They are the small ones, the
    non-US ones, and the ones announced by filing or by an investor's blog post
    rather than by an article.

    Args:
        api_key: Unused, kept for call-site compatibility (the Anthropic client
                 reads ANTHROPIC_API_KEY from the environment)
        start_date: YYYY-MM-DD start date
        end_date: YYYY-MM-DD end date
        topic: 'ai' (default) or 'deeptech'
        passes: source-pass keys to run (see FUNDING_PASS_KEYS). None runs all
                of them, which is days × 5 search turns — the cost of not
                missing rounds. Pass a subset to trade coverage for spend.

    Returns:
        List of funding event dicts with keys: date, company, summary, stage, raise, valuation, investors
    """
    from datetime import timedelta

    start_dt = datetime.strptime(start_date, "%Y-%m-%d")
    end_dt = datetime.strptime(end_date, "%Y-%m-%d")

    total_days = (end_dt - start_dt).days + 1
    days = [(start_dt + timedelta(days=n)).strftime("%Y-%m-%d")
            for n in range(total_days)]

    selected = FUNDING_SOURCE_PASSES
    if passes:
        wanted = {p.strip().lower() for p in passes}
        selected = [p for p in FUNDING_SOURCE_PASSES if p['key'] in wanted]
        unknown = wanted - set(FUNDING_PASS_KEYS)
        if unknown:
            print(f"  WARNING: unknown funding pass(es) {sorted(unknown)} — "
                  f"known passes are {FUNDING_PASS_KEYS}")
        if not selected:
            selected = FUNDING_SOURCE_PASSES

    # Every (day, source pass) combination is one search turn, and they are all
    # independent — so the whole grid goes through parallel_map at once rather
    # than a day's passes waiting on each other.
    tasks = [(day, source_pass) for day in days for source_pass in selected]
    print(f"  Searching {total_days} day(s) × {len(selected)} source pass(es) "
          f"= {len(tasks)} search turns: {start_date} to {end_date} "
          f"({MAX_WORKERS} at a time)")
    print(f"    Passes: {', '.join(p['label'] for p in selected)}")

    done = [0]

    def _report(_i, task, events):
        done[0] += 1
        day, source_pass = task
        print(f"  [{done[0]}/{len(tasks)}] {day} [{source_pass['key']}]: "
              f"{len(events or [])} event(s)")

    per_task = parallel_map(
        lambda task: _search_funding_single_day(api_key, task[0], topic, task[1]),
        tasks, label='funding search', on_result=_report)

    all_events = []
    for events in per_task:
        all_events.extend(events or [])

    unique = _merge_funding_events(all_events, [])
    print(f"  {len(all_events)} result(s) across all passes → "
          f"{len(unique)} unique funding event(s)")

    # Cross-check: how many independent kinds of source found each round. One is
    # not a reason to drop it — a Form D nobody wrote about is still a round —
    # but a table where most rows are single-source is a signal that the passes
    # are not overlapping, which is how the misses happened last time.
    single = [e for e in unique if len(e.get('_passes') or []) < 2]
    if unique:
        print(f"  Corroboration: {len(unique) - len(single)}/{len(unique)} "
              f"confirmed by 2+ source kinds")
        if single:
            names = ', '.join(e.get('company', '?') for e in single[:8])
            print(f"    Single-source: {names}"
                  + (' …' if len(single) > 8 else ''))

    unique = [_prefer_primary_source(e) for e in unique]

    # Drop events outside the requested date range — a web search often returns
    # historical results regardless of the date specified.
    try:
        start_dt = datetime.strptime(start_date, "%Y-%m-%d")
        end_dt = datetime.strptime(end_date, "%Y-%m-%d")
        in_range = []
        out_of_range = 0
        for e in unique:
            event_date_str = e.get('date', '')[:10]
            try:
                event_dt = datetime.strptime(event_date_str, "%Y-%m-%d")
                if start_dt <= event_dt <= end_dt:
                    in_range.append(e)
                else:
                    out_of_range += 1
            except ValueError:
                in_range.append(e)  # keep events with unparseable dates
        if out_of_range:
            print(f"  Removed {out_of_range} event(s) outside {start_date}–{end_date}")
        unique = in_range
    except Exception as e:
        print(f"  WARNING: Date filtering failed ({e}), keeping all events")

    # For AI topic, validate each event is genuinely AI-focused
    if topic == 'ai':
        print(f"  Validating AI relevance...")
        unique = _filter_ai_funding_events(unique)

    return unique


def generate_executive_summary(articles: list) -> str:
    """
    Generate a 2-3 paragraph Chinese executive summary from the week's top articles.

    Uses the highest-relevance articles as context so the model focuses on what
    actually matters for investment decisions, not just volume of news.
    Returns empty string on failure so callers can skip the section gracefully.
    """
    # Feed only the most relevant articles to keep the prompt tight and focused
    top = sorted(articles, key=lambda a: a.get('relevance', 3), reverse=True)[:15]
    if not top:
        return ''

    news_lines = '\n'.join(
        f"- [{a.get('relevance', 3)}/5] {a.get('title', '')}：{a.get('summary', '')[:120]}"
        for a in top
    )

    prompt = (
        "你是一位专注AI领域的风险投资分析师。基于以下本周AI行业重要新闻，生成一份简洁的执行摘要。\n\n"
        "要求：\n"
        "- 共2-3段，每段2-3句\n"
        "- 第一段：本周最重要的技术或产品动态\n"
        "- 第二段：值得关注的融资、并购或公司战略动向\n"
        "- 第三段（可选）：对AI投资格局的整体判断或近期需重点关注的方向\n"
        "- 语言简洁专业，直接切入重点，不要列表\n\n"
        "【事实约束】只能使用下面列出的新闻条目中的信息。不得加入列表之外的公司、"
        "数字、融资金额或行业背景，也不要对未来做预测。列表中信息不足时就写短一些。\n\n"
        f"本周新闻（按投资价值排序）：\n{news_lines}\n\n"
        "只输出摘要正文，不要标题或其他说明。"
    )

    return call_claude(prompt, max_tokens=4096, effort='medium')


def add_executive_summary_section(doc, summary_text: str):
    """
    Write the executive summary as a shaded box at the top of the document.
    """
    h = doc.add_heading('本周执行摘要', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_font(h)

    # Light blue shaded paragraph block
    for para_text in summary_text.split('\n'):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.15)

        # Add subtle left-border shading via paragraph XML
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:space'), '6')
        left.set(qn('w:color'), '4472C4')
        pBdr.append(left)
        pPr.append(pBdr)

        run = p.add_run(para_text)
        set_run_font(run, font_size=10)

    doc.add_paragraph('')


# ── Report header ─────────────────────────────────────────────────────────────

REPORT_TITLE = 'AI 新闻简报'


def _format_range(start_date: str, end_date: str) -> str:
    """'2026年8月10日 – 2026年8月14日' from two ISO dates.

    Falls back to the raw strings if either fails to parse, so a malformed
    --start_date shows up in the header rather than crashing the run.
    """
    try:
        start = datetime.strptime(start_date, '%Y-%m-%d')
        end = datetime.strptime(end_date, '%Y-%m-%d')
    except ValueError:
        return f'{start_date} – {end_date}'

    fmt = '{d.year}年{m}月{day}日'
    return (fmt.format(d=start, m=start.month, day=start.day) + ' – '
            + fmt.format(d=end, m=end.month, day=end.day))


def add_report_header(doc: Document, start_date: str, end_date: str,
                      title: str = None):
    """
    The masthead: a title and the period it covers, and nothing else.

    Deliberately free of the run's bookkeeping — article counts, generation
    timestamps, source tallies. Those are facts about the pipeline, and a reader
    opening a news digest should meet the news, not its build log.

    Shared by both AI report generators so the two layouts open the same way.
    """
    title_para = doc.add_heading(title or REPORT_TITLE, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_font(title_para)

    subtitle = doc.add_paragraph(_format_range(start_date, end_date))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        set_run_font(run, font_size=12)
        run.font.color.rgb = RGBColor(110, 110, 110)

    doc.add_paragraph('')


def create_funding_table(doc: Document, funding_events: list, heading: str = 'AI 融资动态'):
    """
    Add the fundraising section: a 7-column table, and nothing under it.

    The old 优先级 column is gone. It was a stage lookup dressed as editorial
    judgement — "seed → 重点关注" is a rule the reader can apply from the 轮次
    column they are already reading, and printing it as a verdict gave a lookup
    table the authority of a recommendation.

    Args:
        doc: Document object
        funding_events: List of funding event dicts
        heading: Section heading text
    """
    doc.add_paragraph('')
    h = doc.add_heading(heading, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_font(h)

    if not funding_events:
        _set_para_font(doc.add_paragraph('该时间段内未发现融资事件。'))
        return

    # Column widths reallocate the 0.7" the priority column used to take: the
    # summary carries the source links under it and was the tightest cell.
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'

    col_widths = [Inches(0.65), Inches(0.9), Inches(1.95), Inches(0.6),
                  Inches(0.65), Inches(0.65), Inches(1.1)]
    for i, width in enumerate(col_widths):
        table.columns[i].width = width
    _set_table_cell_margins(table)

    # Header row (Chinese)
    headers = ['日期', '公司', '概述', '轮次', '融资额', '估值', '投资方']
    header_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        header_cells[i].text = h_text
        for run in header_cells[i].paragraphs[0].runs:
            run.bold = True
            set_run_font(run, font_size=10)

    # Chronological by announcement: date first, then the announcement time when
    # a source gave one, so several rounds announced on the same day appear in
    # the order they broke rather than in search-result order.
    funding_events.sort(key=lambda x: (x.get('date', ''),
                                       x.get('announced_time', '') or '99:99'))

    # Data rows
    for event in funding_events:
        row_cells = table.add_row().cells

        # Col 0: date
        date_run = row_cells[0].paragraphs[0].add_run(event.get('date', '')[:10])
        set_run_font(date_run, font_size=9)

        # Col 1: company name, hyperlinked if URL available
        company = event.get('company', 'Not disclosed')
        url = event.get('url') or event.get('_url') or ''
        if not url:
            merged = _merge_sources(event)
            url = merged[0] if merged else ''
        company_para = row_cells[1].paragraphs[0]
        if url:
            add_hyperlink(company_para, url, company, font_size=9)
        else:
            run = company_para.add_run(company)
            set_run_font(run, font_size=9)

        # Col 2: Chinese summary, followed by the corroborating source links.
        # The sources belong next to the claim they support rather than in a
        # column of their own — an eighth column on letter paper leaves each one
        # too narrow to read.
        summary_para = row_cells[2].paragraphs[0]
        summary_run = summary_para.add_run(str(event.get('summary', 'Not disclosed')))
        set_run_font(summary_run, font_size=9)

        # A conflict between sources is reported in the row it belongs to, not
        # in a block under the table. Silently picking one figure would present
        # a contested number as settled fact, and a footnote at the bottom of
        # the section is read by nobody looking at the row it qualifies.
        note = (event.get('discrepancy') or '').strip()
        if note:
            note_para = row_cells[2].add_paragraph()
            note_para.paragraph_format.space_before = Pt(2)
            flag_run = note_para.add_run('来源存在出入：')
            flag_run.bold = True
            set_run_font(flag_run, font_size=8)
            flag_run.font.color.rgb = RGBColor(89, 89, 89)
            note_run = note_para.add_run(note)
            set_run_font(note_run, font_size=8)
            note_run.font.color.rgb = RGBColor(89, 89, 89)

        sources = _merge_sources(event)
        if sources:
            src_para = row_cells[2].add_paragraph()
            src_para.paragraph_format.space_before = Pt(2)
            label_run = src_para.add_run('来源 ')
            set_run_font(label_run, font_size=8)
            label_run.font.color.rgb = RGBColor(120, 120, 120)
            for n, src in enumerate(sources[:3], 1):
                if n > 1:
                    sep = src_para.add_run(' · ')
                    set_run_font(sep, font_size=8)
                    sep.font.color.rgb = RGBColor(120, 120, 120)
                add_hyperlink(src_para, src, f'[{n}] {_source_label(src)}', font_size=8)

        # Cols 3-6: remaining fields
        remaining = [
            event.get('stage', 'Not disclosed'),
            event.get('raise', 'Not disclosed'),
            event.get('valuation', 'Not disclosed'),
            event.get('investors', 'Not disclosed'),
        ]
        for i, val in enumerate(remaining, start=3):
            para = row_cells[i].paragraphs[0]
            run = para.add_run(str(val))
            set_run_font(run, font_size=9)


def _source_label(url: str) -> str:
    """Publisher name for a source link — 'techcrunch.com' from a full URL."""
    try:
        host = (urlsplit(url).hostname or '').lower()
    except ValueError:
        return 'source'
    return host[4:] if host.startswith('www.') else (host or 'source')


def convert_bullets_to_paragraph(text: str) -> str:
    """
    Convert bullet point text to paragraph format.
    Removes bullet markers and joins into flowing text.

    Args:
        text: Text with potential bullet points

    Returns:
        Text as paragraph without bullets
    """
    # Remove common bullet markers and clean up
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        line = line.strip()
        # Remove bullet markers
        line = re.sub(r'^[\-\*•]\s*', '', line)
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        if line:
            cleaned_lines.append(line)

    # Join with spaces to form paragraph
    return ' '.join(cleaned_lines)


# ── AI news category grouping ──────────────────────────────────────────────────

AI_NEWS_CATEGORY_ORDER = ["模型与研究", "产品与应用", "大科技公司", "政策与安全", "行业动态", "其他"]

AI_NEWS_CATEGORY_COLORS = {
    "模型与研究": "DDEEFF",
    "产品与应用": "FFF8E1",
    "大科技公司": "E6F4EA",
    "政策与安全": "FCE8E8",
    "行业动态": "F3E5F5",
    "其他":      "F5F5F5",
}

# ── Deeptech category grouping ─────────────────────────────────────────────────

DEEPTECH_CATEGORY_ORDER = ["半导体", "机器人", "新能源", "其他"]

DEEPTECH_CATEGORY_COLORS = {
    "半导体": "E8F0FE",
    "机器人": "E6F4EA",
    "新能源": "FFF8E1",
    "其他":   "F3E5F5",
}

SEMICONDUCTOR_KEYWORDS = [
    "semiconductor", "chip", "芯片", "半导体", "wafer", "fab", "foundry",
    "transistor", "lithography", "eda", "photonic", "asic", "fpga",
    "nvidia", "intel", "amd", "tsmc", "arm ", "risc", "memory", "dram",
    "nand", "soc", "gpu", "cpu", "mpu", "integrated circuit",
]
ROBOTICS_KEYWORDS = [
    "robot", "机器人", "humanoid", "autonomous", "drone", "无人机",
    "unmanned", "exoskeleton", "cobots", "manipulation", "locomotion",
    "actuator", "servo", "mechatronics", "automation",
]
ENERGY_KEYWORDS = [
    "energy", "新能源", "electric vehicle", "ev ", " ev\n", "solar",
    "battery", "电池", "储能", "充电", "wind", "nuclear", "hydrogen",
    "fuel cell", "grid", "power", "renewable", "carbon", "climate",
    "clean tech", "cleantech", "charging", "inverter", "photovoltaic",
]


def classify_deeptech_article(article: dict) -> str:
    """Classify a deeptech article into 半导体 | 机器人 | 新能源 | 其他."""
    text = " " + " ".join([
        article.get("title", ""),
        article.get("summary", ""),
        article.get("description", ""),
    ]).lower() + " "

    for kw in SEMICONDUCTOR_KEYWORDS:
        if kw in text:
            return "半导体"
    for kw in ROBOTICS_KEYWORDS:
        if kw in text:
            return "机器人"
    for kw in ENERGY_KEYWORDS:
        if kw in text:
            return "新能源"
    return "其他"


def _add_deeptech_header_row(table, label: str, fill_hex: str, ncols: int = 2):
    """Add a full-width merged header row for category sections."""
    row = table.add_row()
    row.cells[0].merge(row.cells[ncols - 1])
    cell = row.cells[0]

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

    # Left accent border
    tcBdr = OxmlElement('w:tcBdr')
    left_bdr = OxmlElement('w:left')
    left_bdr.set(qn('w:val'), 'single')
    left_bdr.set(qn('w:sz'), '18')
    left_bdr.set(qn('w:space'), '0')
    left_bdr.set(qn('w:color'), '1F497D')
    tcBdr.append(left_bdr)
    tcPr.append(tcBdr)

    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(label)
    run.bold = True
    set_run_font(run, font_size=11)
    run.font.color.rgb = RGBColor(31, 73, 125)


def create_grouped_deeptech_table(
    doc: Document,
    articles: list,
    chinese_only: bool = False,
    translate: bool = False,
    heading: str = '深科技新闻摘要',
):
    """Build a grouped deeptech news table: 半导体 → 机器人 → 新能源 → 其他."""
    groups = {cat: [] for cat in DEEPTECH_CATEGORY_ORDER}
    for article in articles:
        groups[classify_deeptech_article(article)].append(article)

    for cat in DEEPTECH_CATEGORY_ORDER:
        groups[cat].sort(key=lambda x: x.get("published_at", ""))

    total = sum(len(v) for v in groups.values())

    h = doc.add_heading(heading, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_font(h)
    print(f"  News section: {total} item(s)")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Inches(0.85)
    table.columns[1].width = Inches(5.65)
    _set_table_cell_margins(table)

    hdr = table.rows[0].cells
    hdr[0].text = '日期'
    hdr[1].text = '摘要'
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                set_run_font(run, font_size=10)

    for cat in DEEPTECH_CATEGORY_ORDER:
        cat_articles = groups[cat]
        if not cat_articles:
            continue

        _add_deeptech_header_row(table, cat, DEEPTECH_CATEGORY_COLORS[cat])

        for article in cat_articles:
            row_cells = table.add_row().cells

            date_para = row_cells[0].paragraphs[0]
            date_para.paragraph_format.space_before = Pt(4)
            set_run_font(date_para.add_run(format_date_for_display(article.get('published_at', ''))), font_size=9)

            _fill_summary_cell(row_cells[1], article, translate, chinese_only)

    return table


# Short Chinese badge labels appended inline after each article title in the Word doc.
# 'other' is intentionally omitted — no badge is shown for generic news, keeping
# the visual noise low and making the meaningful signals stand out.
VC_SIGNAL_LABELS = {
    'funding':     '[融资]',
    'product':     '[产品]',
    'partnership': '[合作]',
    'hire':        '[人事]',
    'regulatory':  '[监管]',
    'research':    '[研究]',
}

# Each signal type gets a distinct color so a reader can scan by color at a glance.
VC_SIGNAL_COLORS = {
    'funding':     RGBColor(0, 112, 192),   # blue
    'product':     RGBColor(0, 128, 0),     # green
    'partnership': RGBColor(112, 48, 160),  # purple
    'hire':        RGBColor(197, 90, 17),   # orange
    'regulatory':  RGBColor(192, 0, 0),     # red
    'research':    RGBColor(31, 73, 125),   # dark blue
}


def _load_watchlist(watchlist_file: str) -> list:
    """Load company names from watchlist.txt, one per line. Lines starting with # are comments."""
    if not watchlist_file or not os.path.exists(watchlist_file):
        # Silently skip if no file — watchlist is opt-in, not required
        return []
    with open(watchlist_file, 'r', encoding='utf-8') as f:
        companies = [line.strip() for line in f if line.strip() and not line.startswith('#')]
    if companies:
        print(f"Watchlist loaded: {len(companies)} companies ({', '.join(companies[:5])}{'...' if len(companies) > 5 else ''})")
    return companies


def _check_watchlist(article: dict, watchlist: list) -> list:
    """Return watchlist company names found in article title or summary.

    Checks title + summary + description so it catches both original articles
    and TLDR-style blurbs where the company name only appears in the description.
    """
    text = ' '.join([
        article.get('title', ''),
        article.get('summary', ''),
        article.get('description', ''),
    ]).lower()
    return [co for co in watchlist if co.lower() in text]


def create_watchlist_section(doc: Document, articles: list,
                              translate: bool, chinese_only: bool):
    """Add a Watchlist Highlights section at the top of the document.

    Appears before the main AI News table so investment team members can
    immediately see news about companies they're actively tracking.
    """
    h = doc.add_heading('关注公司动态', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_font(h)


    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.columns[0].width = Inches(0.75)
    table.columns[1].width = Inches(0.45)
    table.columns[2].width = Inches(5.3)
    _set_table_cell_margins(table)

    hdr = table.rows[0].cells
    for i, txt in enumerate(['日期', '优先级', '摘要']):
        hdr[i].text = txt
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            set_run_font(run, font_size=10)
        if i == 1:
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    articles_sorted = sorted(articles, key=lambda x: (-x.get('relevance', 3), x.get('published_at', '')))
    for idx, article in enumerate(articles_sorted):
        _add_article_row(table, article, translate, chinese_only,
                         idx, len(articles_sorted), show_priority=True)


# Translations are done up front, in parallel, then read from here while the
# document is built. python-docx is not thread-safe and the build is sequential
# by nature, so translating inside the render loop meant one blocking API call
# per article with nothing else in flight.
_TRANSLATIONS: dict = {}


def _translate_with_retry(text: str) -> str:
    for attempt in range(3):
        out = translate_to_chinese_claude(text)
        if out:
            return out
        time.sleep(2 * (attempt + 1))
    return ''


def pretranslate(texts: list):
    """Translate every distinct string now, concurrently, into _TRANSLATIONS."""
    todo = [t for t in dict.fromkeys(texts) if t and t not in _TRANSLATIONS]
    if not todo:
        return
    print(f"  Translating {len(todo)} text(s) ({MAX_WORKERS} at a time)...")
    results = parallel_map(_translate_with_retry, todo, label='translate')
    for text, result in zip(todo, results):
        # Falling back to the original text keeps the report readable when a
        # translation fails, which is what the old inline retry loop did too.
        _TRANSLATIONS[text] = result or text


def translate_cached(text: str) -> str:
    """The pre-translated string, translating on the spot if it was missed."""
    if not text:
        return text
    if text not in _TRANSLATIONS:
        _TRANSLATIONS[text] = _translate_with_retry(text) or text
    return _TRANSLATIONS[text]


def _fill_summary_cell(cell, article: dict, translate: bool, chinese_only: bool):
    """Fill the summary cell: hyperlinked title + body paragraph(s) with proper spacing."""
    # Headlines collected from X arrive as the tweet — handle, hashtags, thread
    # marker and the ellipsis where the post ran out of characters.
    title = clean_headline(article.get('title', '')) or '无标题'
    # The news behind the post, never the post itself — an x.com link is
    # provenance, not a source. Falls back to an unlinked headline.
    url = news_link(article)
    vc_signal = article.get('vc_signal', '')

    title_para = cell.paragraphs[0]
    title_para.paragraph_format.space_after = Pt(3)
    if url:
        add_hyperlink(title_para, url, title, font_size=10)
    else:
        run = title_para.add_run(title)
        run.bold = True
        set_run_font(run, font_size=10)

    # Append a colored VC signal badge inline with the title (e.g. "  [融资]" in blue).
    # 'other' is skipped — only named signal types get a badge to avoid visual clutter.
    if vc_signal and vc_signal != 'other' and vc_signal in VC_SIGNAL_LABELS:
        badge_run = title_para.add_run('  ' + VC_SIGNAL_LABELS[vc_signal])
        badge_run.bold = True
        set_run_font(badge_run, font_size=9)
        badge_run.font.color.rgb = VC_SIGNAL_COLORS[vc_signal]

    # `or` rather than a .get default: a summarization that failed leaves an
    # empty string, and a present-but-empty key would render a blank cell.
    summary = (article.get('summary') or article.get('description')
               or article.get('title') or '暂无摘要')
    summary_text = convert_bullets_to_paragraph(summary)

    if chinese_only:
        body = cell.add_paragraph()
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(4)
        add_formatted_text(body, summary_text, font_size=10)
    elif translate:
        body = cell.add_paragraph()
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(4)
        add_formatted_text(body, translate_cached(summary_text), font_size=10)
    else:
        body = cell.add_paragraph()
        body.paragraph_format.space_before = Pt(0)
        body.paragraph_format.space_after = Pt(4)
        add_formatted_text(body, summary_text, font_size=10)


_PRIORITY_COLORS = {
    5: RGBColor(31, 73, 125),   # dark blue
    4: RGBColor(68, 114, 196),  # medium blue
    3: RGBColor(89, 89, 89),    # dark gray
    2: RGBColor(128, 128, 128), # gray
    1: RGBColor(166, 166, 166), # light gray
}


def _add_article_row(table, article: dict, translate: bool,
                     chinese_only: bool, idx: int, total: int,
                     show_priority: bool = False):
    """Write one article as a table row. Shared by flat and grouped layouts."""
    row_cells = table.add_row().cells

    date_para = row_cells[0].paragraphs[0]
    date_para.paragraph_format.space_before = Pt(4)
    date_run = date_para.add_run(format_date_for_display(article.get('published_at', '')))
    set_run_font(date_run, font_size=9)

    if show_priority:
        relevance = article.get('relevance', 3)
        p_para = row_cells[1].paragraphs[0]
        p_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_para.paragraph_format.space_before = Pt(4)
        p_run = p_para.add_run(str(relevance))
        p_run.bold = True
        set_run_font(p_run, font_size=13)
        p_run.font.color.rgb = _PRIORITY_COLORS.get(relevance, RGBColor(89, 89, 89))
        _fill_summary_cell(row_cells[2], article, translate, chinese_only)
    else:
        _fill_summary_cell(row_cells[1], article, translate, chinese_only)


def create_news_table(doc: Document, articles: list, max_articles: int = None,
                      translate: bool = False,
                      chinese_only: bool = False):
    """Create AI News table. Groups by category + shows priority column when articles are categorized."""
    if max_articles:
        articles = articles[:max_articles]

    heading = doc.add_heading('AI 新闻摘要', level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para_font(heading)
    # No "共 N 篇" line: how many items the pipeline produced is a fact about the
    # pipeline, not news. The count still goes to the console during the run.
    print(f"  News section: {len(articles)} item(s)")

    is_categorized = any('category' in a for a in articles)

    if is_categorized:
        # 3-column table: Date | 优先级 | Summary
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.columns[0].width = Inches(0.75)
        table.columns[1].width = Inches(0.45)
        table.columns[2].width = Inches(5.3)
        _set_table_cell_margins(table)

        hdr = table.rows[0].cells
        for i, txt in enumerate(['日期', '优先级', '摘要']):
            hdr[i].text = txt
            for run in hdr[i].paragraphs[0].runs:
                run.bold = True
                set_run_font(run, font_size=10)
            if i == 1:
                hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # 2-column flat table: Date | Summary
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        table.columns[0].width = Inches(0.85)
        table.columns[1].width = Inches(5.65)
        _set_table_cell_margins(table)

        hdr = table.rows[0].cells
        hdr[0].text = '日期'
        hdr[1].text = '摘要'
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    set_run_font(run, font_size=10)

    if is_categorized:
        groups: dict = {cat: [] for cat in AI_NEWS_CATEGORY_ORDER}
        for article in articles:
            cat = article.get('category', '其他')
            if cat not in groups:
                cat = '其他'
            groups[cat].append(article)

        idx = 0
        for cat in AI_NEWS_CATEGORY_ORDER:
            cat_articles = groups[cat]
            if not cat_articles:
                continue
            cat_articles.sort(key=lambda x: (-x.get('relevance', 3), x.get('published_at', '')))
            _add_deeptech_header_row(table, cat, AI_NEWS_CATEGORY_COLORS[cat], ncols=3)
            for article in cat_articles:
                _add_article_row(table, article, translate, chinese_only,
                                 idx, len(articles), show_priority=True)
                idx += 1
    else:
        articles.sort(key=lambda x: x.get('published_at', ''))
        for i, article in enumerate(articles):
            _add_article_row(table, article, translate, chinese_only,
                             i, len(articles), show_priority=False)

    return table


def generate_word_doc(start_date: str, end_date: str,
                      articles_file: str = '.tmp/summarized_articles.json',
                      output_dir: str = 'output',
                      max_articles: int = None,
                      translate: bool = False,
                      chinese_only: bool = False,
                      output_prefix: str = 'AI_News',
                      funding_topic: str = 'ai',
                      doc_title: str = None,
                      min_signal: int = 3,
                      watchlist_file: str = 'watchlist.txt',
                      include_frontier: bool = False,
                      include_opinion: bool = False,
                      require_source: bool = True,
                      funding_passes: list = None,
                      watchlist_section: bool = False):
    """
    Main document generation function

    Args:
        start_date: YYYY-MM-DD
        end_date: YYYY-MM-DD
        articles_file: Path to summarized articles JSON
        output_dir: Output directory
        max_articles: Maximum articles to include (None = all)
        translate: Add Chinese translation using Claude
    """
    load_dotenv(override=True)

    print("Loading data...")

    # Load summarized articles
    if not os.path.exists(articles_file):
        print(f"ERROR: {articles_file} not found")
        print("Make sure to run summarize_articles.py first")
        sys.exit(1)

    with open(articles_file, 'r', encoding='utf-8') as f:
        articles = json.load(f)

    print(f"Loaded {len(articles)} articles")

    # Editorial rules: no duplicates, and for the AI brief no frontier labs and
    # no opinion pieces. The deeptech report shares the dedup but not the
    # startups-only line — that is the AI report's editorial choice, not a
    # property of every report this function builds.
    print("Curating...")
    articles = curate(articles, include_frontier, include_opinion,
                      dedupe_only=(funding_topic != 'ai'),
                      require_source=require_source)

    # The AI news section carries AI stories only, and never a funding round —
    # rounds are the fundraising table's subject, with better columns and wider
    # sourcing. This also applies the relevance floor (--min-signal), so the
    # tally names one reason per article instead of two filters reporting
    # separately. See filter_news_section() in news_filters.py.
    #
    # The deeptech report is out of scope: "no funding rounds, must be AI" is
    # the AI brief's editorial line, not a property of every report built here.
    if funding_topic == 'ai':
        articles, _off_topic = filter_news_section(articles, min_signal=min_signal)
    elif min_signal > 1:
        before = len(articles)
        articles = [a for a in articles if signal_score(a) >= min_signal]
        print(f"Signal filter (>= {min_signal}): keeping {len(articles)}/{before} articles")

    # Load the watchlist of investment-target company names (edit watchlist.txt to configure).
    # Returns an empty list if the file doesn't exist, so the watchlist section is skipped.
    watchlist = _load_watchlist(watchlist_file)

    if translate:
        print("Translation enabled (Claude)")
        # Every summary that will be rendered, translated in one concurrent
        # batch before the document is built. Doing it here also means the
        # per-article "[3/40] Translating..." wait is gone from the render loop.
        pretranslate([convert_bullets_to_paragraph(
            a.get('summary', a.get('description', ''))) for a in articles])

    # Check for the Anthropic key for funding extraction
    anthropic_key = os.getenv('ANTHROPIC_API_KEY')

    # Create document
    print("Creating Word document...")
    doc = Document()

    # Set 1" margins — gives 6.5" content width on letter paper
    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # Header: title and date range, nothing else. The generation timestamp and
    # 文章总数 line that used to sit here were build metadata printed as if they
    # were the lead.
    add_report_header(doc, start_date, end_date, doc_title)

    # Watchlist: the matches are reported on the console, not as a section in the
    # document. The AI report's structure is header → news → fundraising and
    # nothing else, and a second news table above the first one competing for the
    # top of the page is the thing that rule exists to prevent. Pass
    # --watchlist-section to put it back.
    if watchlist:
        watchlist_hits = []
        for article in articles:
            matches = _check_watchlist(article, watchlist)
            if matches:
                article['_watchlist_matches'] = matches
                watchlist_hits.append(article)
        if watchlist_hits:
            names = sorted({co for a in watchlist_hits
                            for co in a['_watchlist_matches']})
            print(f"Watchlist: {len(watchlist_hits)} article(s) matched "
                  f"({', '.join(names)})")
            if watchlist_section:
                create_watchlist_section(doc, watchlist_hits, translate, chinese_only)
                _add_horizontal_rule(doc)
        else:
            print("Watchlist: no matches found this period")

    # Create news table (grouped for deeptech, flat for others)
    print("Creating news table...")
    if funding_topic == 'deeptech':
        create_grouped_deeptech_table(doc, articles, chinese_only, translate)
    else:
        create_news_table(doc, articles, max_articles, translate, chinese_only)

    # Two-source funding pipeline, both on Claude:
    #   1. Extraction from the articles we already collected — high fidelity,
    #      because these articles were hand-curated by our sources.
    #   2. Web search supplements with any funding events that weren't covered
    #      by TechCrunch / TLDR (smaller raises, non-English coverage, etc.).
    # Both lists are merged by _merge_funding_events(), keeping the richer entry per company.
    # Article extraction is AI-only for now; deeptech relies solely on web search.
    topic_label = "Deeptech" if funding_topic == "deeptech" else "AI"
    all_funding_events = []

    if funding_topic == 'ai':
        print("Extracting funding events from collected articles (Claude)...")
        article_events = extract_funding_with_claude(articles)
        all_funding_events.extend(article_events)

    if anthropic_key:
        print(f"Supplementing with {topic_label} funding web search (Claude)...")
        web_events = extract_funding_with_web_search(
            anthropic_key, start_date, end_date, funding_topic, funding_passes)
        all_funding_events = _merge_funding_events(all_funding_events, web_events)
        print(f"  Total after merge: {len(all_funding_events)} funding events")
    elif not all_funding_events:
        print("  WARNING: ANTHROPIC_API_KEY not set and no extracted events found, skipping funding section")

    if all_funding_events:
        emit_funding_claims(all_funding_events)
        heading_map = {"AI": "AI 融资动态", "Deeptech": "深科技融资动态"}
        create_funding_table(doc, all_funding_events, heading=heading_map.get(topic_label, f"{topic_label} 融资动态"))

    # Save document
    os.makedirs(output_dir, exist_ok=True)
    filename = f'{output_prefix}_{start_date.replace("-", "")}_{end_date.replace("-", "")}.docx'
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)

    print(f"\n✓ Document saved to {filepath}")
    return filepath


def main():
    """Main execution function"""
    parser = argparse.ArgumentParser(description='Generate AI News Word document')
    parser.add_argument('--start_date', required=True, help='Start date (YYYY-MM-DD)')
    parser.add_argument('--end_date', required=True, help='End date (YYYY-MM-DD)')
    parser.add_argument('--articles', default='.tmp/summarized_articles.json', help='Summarized articles file')
    parser.add_argument('--output_dir', default='output', help='Output directory')
    parser.add_argument('--max', type=int, default=None, help='Maximum articles to include (default: all)')
    parser.add_argument('--translate', action='store_true', help='Add Chinese translation using Claude')
    parser.add_argument('--chinese-only', action='store_true', help='Output Chinese summary only (no English), for pre-summarized Chinese articles')
    parser.add_argument('--output-prefix', default='AI_News', help='Output filename prefix (default: AI_News)')
    parser.add_argument('--funding-topic', choices=['ai', 'deeptech'], default='ai',
                        help='Funding search topic: ai (default) or deeptech')
    parser.add_argument('--doc-title', default=None, help='Document title (default: AI News Report)')
    parser.add_argument('--min-signal', type=int, default=3, metavar='N',
                        help='Minimum relevance score 1-5 to include in news table (default: 1 = all). '
                             'Use 3 to drop low-value articles, 4 for high-signal-only.')
    parser.add_argument('--watchlist', default='watchlist.txt', metavar='FILE',
                        help='Path to watchlist file (default: watchlist.txt). '
                             'One company name per line. Matching articles appear in a highlights section.')
    parser.add_argument('--include-frontier', action='store_true',
                        help='Keep stories about frontier labs and big tech '
                             '(frontier_labs.txt). Off by default for --funding-topic ai: '
                             'the news table covers smaller AI startups.')
    parser.add_argument('--include-opinion', action='store_true',
                        help='Keep commentary, opinion pieces and executive statements. '
                             'Off by default: the news table is for reported events.')
    parser.add_argument('--watchlist-section', action='store_true',
                        help='Render the watchlist companies as their own section '
                             'above the news table. Off by default: the report is '
                             'header, news, fundraising and nothing else. Matches '
                             'are reported on the console either way.')
    parser.add_argument('--funding-passes', default=None, metavar='KEYS',
                        help='Comma-separated funding source passes to run '
                             '(%s). Each is one web search per day, so all of '
                             'them is days x 5 search turns. Default: all — '
                             'fewer passes means missed rounds.'
                             % ','.join(FUNDING_PASS_KEYS))
    parser.add_argument('--allow-unlinked', dest='require_source',
                        action='store_false',
                        help='Keep articles with no link to published coverage, as '
                             'unlinked headlines. By default they are dropped — run '
                             'tools/resolve_sources.py first to find their coverage.')
    args = parser.parse_args()

    generate_word_doc(
        args.start_date,
        args.end_date,
        args.articles,
        args.output_dir,
        args.max,
        args.translate,
        args.chinese_only,
        args.output_prefix,
        args.funding_topic,
        args.doc_title,
        args.min_signal,
        args.watchlist,
        args.include_frontier,
        args.include_opinion,
        args.require_source,
        [k for k in (args.funding_passes or '').split(',') if k.strip()] or None,
        args.watchlist_section,
    )


def _self_test() -> int:
    """
    Consolidation tests: `python tools/generate_word_doc.py --self-test`.

    This is the subtlest logic in the funding path. Merging two reports of one
    round is the whole point of the section, and merging two *different* rounds
    of one company silently deletes a round — so both directions are pinned
    here rather than checked by hand.
    """
    def ev(**kw):
        base = dict(date='', announced_time='', company='', summary='',
                    stage='Not disclosed', raise_='Not disclosed',
                    valuation='Not disclosed', investors='Not disclosed',
                    url='', sources=[], discrepancy='')
        base.update(kw)
        base['raise'] = base.pop('raise_')
        return base

    failures = []

    def expect(label, a, b, want):
        got = len(_merge_funding_events([a], [b]))
        if got != want:
            failures.append(f"{label}: {got} entries, expected {want}")

    expect('recap of one round merges',
           ev(date='2026-08-10', company='Acme AI', stage='Series A', raise_='US$20 million'),
           ev(date='2026-08-12', company='Acme AI', stage='Series A', raise_='US$20 million'), 1)
    expect('different rounds stay separate',
           ev(date='2026-03-01', company='Acme AI', stage='Seed', raise_='US$3 million'),
           ev(date='2026-08-01', company='Acme AI', stage='Series B', raise_='US$60 million'), 2)
    expect('same amount, one report omits stage',
           ev(date='2026-08-10', company='Beta AI', stage='Series B', raise_='US$40 million'),
           ev(date='2026-08-11', company='Beta AI', raise_='US$40 million'), 1)
    expect('same amount written differently',
           ev(date='2026-08-10', company='Gamma AI', stage='Series A', raise_='US$20 million'),
           ev(date='2026-08-10', company='Gamma AI', stage='Series A', raise_='$20M'), 1)
    expect('two stages announced the same day stay separate',
           ev(date='2026-08-10', company='Delta AI', stage='Seed', raise_='US$5 million'),
           ev(date='2026-08-10', company='Delta AI', stage='Series A', raise_='US$30 million'), 2)
    expect('nothing disclosed, days apart, is one round',
           ev(date='2026-08-10', company='Eps AI'), ev(date='2026-08-11', company='Eps AI'), 1)
    expect('nothing disclosed, weeks apart, is two rounds',
           ev(date='2026-08-01', company='Zeta AI'), ev(date='2026-08-22', company='Zeta AI'), 2)

    # A recap must never move the round's date, and must still contribute detail.
    first = ev(date='2026-08-10', announced_time='08:00', company='Acme AI',
               stage='Series A', raise_='US$20 million', url='https://first.example',
               sources=['https://first.example'])
    recap = ev(date='2026-08-12', company='Acme AI', stage='Series A',
               raise_='US$20 million', valuation='US$150 million',
               investors='Benchmark (lead)', url='https://recap.example',
               sources=['https://recap.example'], discrepancy='amount disputed')
    for label, merged in (('forward', _merge_funding_events([first], [recap])[0]),
                          ('reversed', _merge_funding_events([recap], [first])[0])):
        if merged['date'] != '2026-08-10':
            failures.append(f"{label}: recap date won ({merged['date']})")
        if merged['valuation'] != 'US$150 million':
            failures.append(f"{label}: lost the recap's disclosed valuation")
        if merged['announced_time'] != '08:00':
            failures.append(f"{label}: lost the original announcement time")
        if len(_merge_sources(merged)) != 2:
            failures.append(f"{label}: sources not unioned")
        if not merged['discrepancy']:
            failures.append(f"{label}: discrepancy dropped")

    # Both spellings of "unknown" have to read as undisclosed.
    for value, want in (('不详', False), ('Not disclosed', False), ('', False),
                        ('US$5 million', True)):
        if _is_disclosed(value) != want:
            failures.append(f"_is_disclosed({value!r}) should be {want}")

    if failures:
        print(f"✗ {len(failures)} failure(s):")
        for f in failures:
            print(f"    {f}")
        return 1
    print("✓ generate_word_doc: funding consolidation self-tests passed")
    return 0


if __name__ == "__main__":
    if '--self-test' in sys.argv:
        sys.exit(_self_test())
    main()
