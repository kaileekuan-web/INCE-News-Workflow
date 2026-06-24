#!/usr/bin/env python3
"""
Generate Investment Intelligence Word document.

Structure:
  - Cover: title, date range, signal count
  - Per theme (sorted by signal strength desc):
      one-liner insight | supporting articles | bull | bear | partner take
  - Deal Sourcing: top 5 companies to look at this week

Input:  .tmp/debate_memos.json, .tmp/deal_sourcing.json
Output: output/Investment_Intelligence_YYYYMMDD_YYYYMMDD.docx
"""

import os
import sys
import json
import argparse
import time
import requests
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── Colors ─────────────────────────────────────────────────────────────────────
COLOR_BULL    = RGBColor(0x16, 0x7A, 0x47)   # green
COLOR_BEAR    = RGBColor(0xC0, 0x39, 0x2B)   # red
COLOR_PARTNER = RGBColor(0x1A, 0x53, 0xFF)   # blue
COLOR_HEADER  = RGBColor(0x1F, 0x38, 0x64)   # dark navy
COLOR_INSIGHT = RGBColor(0x2C, 0x3E, 0x50)   # dark slate
COLOR_SIGNAL  = [
    RGBColor(0xBD, 0xC3, 0xC7),  # 1 — grey
    RGBColor(0xF3, 0x9C, 0x12),  # 2 — orange
    RGBColor(0xE6, 0x7E, 0x22),  # 3 — darker orange
    RGBColor(0x27, 0xAE, 0x60),  # 4 — green
    RGBColor(0x8E, 0x44, 0xAD),  # 5 — purple (exceptional)
]

SIGNAL_LABELS = {1: "Low", 2: "Watch", 3: "Moderate", 4: "High", 5: "Exceptional"}


def _set_font(run, size=10, bold=False, color=None, font_name="Calibri"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = font_name
    if color:
        run.font.color.rgb = color


def _para(doc, text="", size=10, bold=False, color=None, align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, color=color)
    return p


def _add_hyperlink(paragraph, url: str, text: str, size=10):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _section_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _signal_badge(para, strength: int):
    stars = "★" * strength + "☆" * (5 - strength)
    label = SIGNAL_LABELS.get(strength, "")
    color = COLOR_SIGNAL[strength - 1] if 1 <= strength <= 5 else COLOR_SIGNAL[0]
    run = para.add_run(f"  {stars} {label} ({strength}/5)")
    _set_font(run, size=9, bold=True, color=color)


def add_theme_section(doc, memo: dict, idx: int, zh: bool = False):
    theme = memo.get("theme", "")
    insight = memo.get("insight", "")
    signal = memo.get("signal_strength", 0)
    rationale = memo.get("rationale", "")
    articles = memo.get("articles", [])
    debate = memo.get("debate", {})

    # Theme header
    p = _para(doc, f"{idx}. {theme}", size=13, bold=True, color=COLOR_HEADER,
              space_before=10, space_after=2)
    _signal_badge(p, signal)

    # One-liner insight (callout style)
    insight_para = doc.add_paragraph()
    insight_para.paragraph_format.space_before = Pt(2)
    insight_para.paragraph_format.space_after = Pt(6)
    insight_para.paragraph_format.left_indent = Inches(0.2)
    run = insight_para.add_run(f'"{insight}"')
    _set_font(run, size=10, bold=True, color=COLOR_INSIGHT)

    # Rationale
    if rationale:
        _para(doc, rationale, size=9.5, color=RGBColor(0x55, 0x55, 0x55),
              space_before=0, space_after=6)

    # Supporting articles
    if articles:
        articles_label = LABELS_ZH["Supporting Articles"] if zh else "Supporting Articles"
        p_label = _para(doc, articles_label, size=9, bold=True,
                        color=RGBColor(0x88, 0x88, 0x88), space_before=0, space_after=2)
        for a in articles:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_before = Pt(0)
            bullet.paragraph_format.space_after = Pt(1)
            if a.get("url"):
                _add_hyperlink(bullet, a["url"], a.get("title", a["url"]), size=9)
                src_run = bullet.add_run(f"  ({a.get('source', '')})")
                _set_font(src_run, size=8.5, color=RGBColor(0xAA, 0xAA, 0xAA))
            else:
                run = bullet.add_run(a.get("title", ""))
                _set_font(run, size=9)

    # Bull case
    bull_label   = LABELS_ZH["Bull Case"]      if zh else "Bull Case"
    bear_label   = LABELS_ZH["Bear Case"]      if zh else "Bear Case"
    partner_label = LABELS_ZH["Partner's Take"] if zh else "Partner's Take"

    if debate.get("bull"):
        _para(doc, bull_label, size=10, bold=True, color=COLOR_BULL,
              space_before=8, space_after=2)
        _para(doc, debate["bull"], size=9.5, space_before=0, space_after=6)

    if debate.get("bear"):
        _para(doc, bear_label, size=10, bold=True, color=COLOR_BEAR,
              space_before=0, space_after=2)
        _para(doc, debate["bear"], size=9.5, space_before=0, space_after=6)

    if debate.get("partner"):
        _para(doc, partner_label, size=10, bold=True, color=COLOR_PARTNER,
              space_before=0, space_after=2)
        _para(doc, debate["partner"], size=9.5, space_before=0, space_after=6)

    _section_divider(doc)


def add_deal_sourcing_section(doc, deals: list, zh: bool = False):
    if not deals:
        return

    doc.add_page_break()
    section_title = LABELS_ZH["Deal Sourcing: Top Companies to Look at This Week"] if zh else \
                    "Deal Sourcing: Top Companies to Look at This Week"
    _para(doc, section_title, size=14, bold=True, color=COLOR_HEADER, space_before=0, space_after=8)

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = DEAL_HEADERS_ZH if zh else ["Company", "Stage", "Theme Fit", "Why Now", "Raise"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        _set_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F3864")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    for deal in deals:
        row = table.add_row().cells
        # Company (hyperlinked)
        p = row[0].paragraphs[0]
        url = deal.get("url", "")
        name = deal.get("company", "")
        if url:
            _add_hyperlink(p, url, name, size=9)
        else:
            r = p.add_run(name)
            _set_font(r, size=9)

        for cell, key in zip(row[1:], ["stage", "thesis_fit", "why_now", "raise"]):
            r = cell.paragraphs[0].add_run(deal.get(key) or "—")
            _set_font(r, size=9)

    _para(doc, "", space_before=6, space_after=0)


HOW_TO_ZH = [
    ("本报告从本周新闻中识别投资主题，并通过结构化三方辩论对每个主题进行压力测试，最后推荐值得关注的公司。"),
    "",
    ("每个主题包含三个视角："),
    ("  多方观点（绿色） — 为何 INCE 应关注并考虑投资的最强论据。聚焦市场机会、时机判断与竞争壁垒。"),
    ("  空方观点（红色） — 对多方论点的强力反驳。揭示风险、市场拥挤以及信号可能为噪音的原因。"),
    ("  合伙人结论（蓝色） — 综合两方观点给出最终判断，并提出 INCE 本周的具体行动建议。"),
    "",
    ("为何采用三个独立智能体？单一提示词评估投资机会时默认给出平衡答案——有用但缺乏锐度。"
     "强制设定对立立场能产生真实的观点张力，从而揭示单一回答会平滑掉的风险与机遇。"),
]

LABELS_ZH = {
    "Bull Case":       "多方观点",
    "Bear Case":       "空方观点",
    "Partner's Take":  "合伙人结论",
    "How to Read This Report": "如何阅读本报告",
    "Supporting Articles":     "相关文章",
    "Deal Sourcing: Top Companies to Look at This Week": "本周重点关注公司",
    "Investment Intelligence Report": "投资洞察报告",
}

DEAL_HEADERS_ZH = ["公司", "轮次", "主题契合度", "关注原因", "融资额"]


def _translate_single(text: str, api_key: str) -> str:
    """Translate a single text string to Chinese. Returns original on failure."""
    if not text:
        return text
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "Content-Type": "application/json",
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
    }
    payload = {
        "model": "claude-haiku-4-5-20251001",
        "max_tokens": 4096,
        "messages": [{"role": "user", "content":
            "将以下文本翻译成简体中文。公司名、人名保持英文。只输出翻译结果，不要其他说明。\n\n" + text}],
    }
    for attempt in range(3):
        try:
            resp = requests.post(url, json=payload, headers=headers, timeout=120)
            resp.raise_for_status()
            result = resp.json()["content"][0]["text"].strip()
            if result:
                return result
        except Exception as e:
            if attempt == 2:
                print(f"  WARNING: Translation failed: {e} — using original")
                return text
        time.sleep(4 * (attempt + 1))
    return text


def _batch_translate(texts: dict, api_key: str) -> dict:
    """Translate a dict of SHORT {key: english_text} fields to Chinese in one API call.
    Only use for short fields (theme, insight, rationale) — NOT long paragraphs."""
    prompt = (
        "将以下 JSON 中所有字符串值翻译成简体中文。"
        "保留 JSON 结构和键名不变。公司名、人名保持英文。只输出 JSON，不要其他内容。\n\n"
        + json.dumps(texts, ensure_ascii=False)
    )
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "Content-Type": "application/json",
        "x-api-key": api_key,
        "anthropic-version": "2023-06-01",
    }
    payload = {
        "model": "claude-haiku-4-5-20251001",
        "max_tokens": 2048,
        "messages": [{"role": "user", "content": prompt}],
    }
    for attempt in range(3):
        try:
            resp = requests.post(url, json=payload, headers=headers, timeout=60)
            resp.raise_for_status()
            text = resp.json()["content"][0]["text"].strip()
            if "```" in text:
                text = text.split("```")[1]
                if text.startswith("json"):
                    text = text[4:]
            return json.loads(text.strip())
        except Exception as e:
            if attempt == 2:
                print(f"  WARNING: Batch translation failed: {e} — using English")
                return texts
            time.sleep(4 * (attempt + 1))
    return texts


def generate_insight_doc(memos_path: str, deals_path: str,
                          start_date: str, end_date: str, output_dir: str,
                          language: str = "en") -> str:
    zh = language == "zh"
    api_key = os.getenv("ANTHROPIC_API_KEY") if zh else None

    with open(memos_path, encoding="utf-8") as f:
        memos = json.load(f)

    # Translate all memo text per theme:
    # - Short fields (theme/insight/rationale) in one batch call
    # - Long fields (bull/bear/partner) individually to avoid JSON truncation
    if zh and api_key:
        print("Translating themes to Chinese...")
        translated_memos = []
        for i, memo in enumerate(memos):
            print(f"  Translating theme {i+1}/{len(memos)}: {memo.get('theme', '')[:60]}")
            debate = memo.get("debate", {})

            # Batch the short fields
            short = _batch_translate({
                "theme":    memo.get("theme", ""),
                "insight":  memo.get("insight", ""),
                "rationale": memo.get("rationale", ""),
            }, api_key)
            time.sleep(1)

            # Translate long debate fields individually
            bull    = _translate_single(debate.get("bull", ""),    api_key); time.sleep(1)
            bear    = _translate_single(debate.get("bear", ""),    api_key); time.sleep(1)
            partner = _translate_single(debate.get("partner", ""), api_key); time.sleep(1)

            m = dict(memo)
            m["theme"]     = short.get("theme",    memo.get("theme", ""))
            m["insight"]   = short.get("insight",  memo.get("insight", ""))
            m["rationale"] = short.get("rationale", memo.get("rationale", ""))
            m["debate"]    = {"bull": bull, "bear": bear, "partner": partner}
            translated_memos.append(m)
        memos = translated_memos

    deals = []
    if deals_path and os.path.exists(deals_path):
        with open(deals_path, encoding="utf-8") as f:
            deals = json.load(f)
        if zh and api_key and deals:
            print("Translating deal sourcing to Chinese...")
            for deal in deals:
                deal["thesis_fit"] = _translate_single(deal.get("thesis_fit", ""), api_key)
                time.sleep(1)
                deal["why_now"] = _translate_single(deal.get("why_now", ""), api_key)
                time.sleep(1)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title_text = LABELS_ZH["Investment Intelligence Report"] if zh else "Investment Intelligence Report"
    _para(doc, title_text, size=20, bold=True,
          color=COLOR_HEADER, align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=0, space_after=4)
    _para(doc, f"{start_date}  →  {end_date}", size=11,
          color=RGBColor(0x77, 0x77, 0x77), align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=0, space_after=2)
    generated_label = "生成时间" if zh else "Generated"
    themes_label = f"{'识别主题数' if zh else 'themes identified'}: {len(memos)}"
    _para(doc, f"{generated_label}: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  {themes_label}",
          size=9, color=RGBColor(0xAA, 0xAA, 0xAA), align=WD_ALIGN_PARAGRAPH.CENTER,
          space_before=0, space_after=10)

    _section_divider(doc)

    # How to read this report
    how_to_title = LABELS_ZH["How to Read This Report"] if zh else "How to Read This Report"
    _para(doc, how_to_title, size=11, bold=True, color=COLOR_HEADER,
          space_before=10, space_after=4)

    how_to_en = [
        ("This report identifies investment themes from the week's news, then stress-tests each "
         "one through a structured three-agent debate before surfacing deal sourcing targets."),
        "",
        ("Each theme section contains three voices:"),
        ("  Bull Case (green) — The strongest argument for why INCE should pay attention and "
         "potentially invest. Focuses on market opportunity, timing, and defensibility."),
        ("  Bear Case (red) — Pushes back hard on the bull thesis. Surfaces risks, crowding, "
         "and reasons the signal might be noise."),
        ("  Partner's Take (blue) — Synthesizes both views into a final verdict and ends with "
         "one specific action item for INCE this week."),
        "",
        ("Why three agents? A single prompt asked to evaluate an opportunity produces a balanced "
         "answer by default — useful but not sharp. Forcing opposing mandates creates real tension "
         "between views, surfacing risks and opportunities a single response would smooth over."),
    ]

    how_to = HOW_TO_ZH if zh else how_to_en

    for line in how_to:
        if line == "":
            _para(doc, "", space_before=0, space_after=3)
        elif line.startswith("  "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            label, _, rest = line.strip().partition(" — ")
            is_bull = "Bull" in label or "多方" in label
            is_bear = "Bear" in label or "空方" in label
            color = COLOR_BULL if is_bull else COLOR_BEAR if is_bear else COLOR_PARTNER
            r1 = p.add_run(label + " — ")
            _set_font(r1, size=9.5, bold=True, color=color)
            r2 = p.add_run(rest)
            _set_font(r2, size=9.5)
        else:
            _para(doc, line, size=9.5, space_before=0, space_after=2)

    _section_divider(doc)
    _para(doc, "", space_before=0, space_after=6)

    # Theme sections
    for i, memo in enumerate(memos, 1):
        add_theme_section(doc, memo, i, zh=zh)

    # Deal sourcing
    add_deal_sourcing_section(doc, deals, zh=zh)

    # Save
    os.makedirs(output_dir, exist_ok=True)
    s = start_date.replace("-", "")
    e = end_date.replace("-", "")
    lang_suffix = "_ZH" if zh else ""
    out_path = os.path.join(output_dir, f"Investment_Intelligence_{s}_{e}{lang_suffix}.docx")
    doc.save(out_path)
    print(f"✓ Saved → {out_path}")
    return out_path


def main():
    parser = argparse.ArgumentParser(description="Generate Investment Intelligence Word doc")
    parser.add_argument("--memos",      required=True, help="debate_memos.json path")
    parser.add_argument("--deals",      default=None,  help="deal_sourcing.json path (optional)")
    parser.add_argument("--start_date", required=True)
    parser.add_argument("--end_date",   required=True)
    parser.add_argument("--output_dir", default="output")
    parser.add_argument("--language",   default="en", choices=["en", "zh"])
    args = parser.parse_args()

    generate_insight_doc(args.memos, args.deals, args.start_date, args.end_date,
                         args.output_dir, language=args.language)


if __name__ == "__main__":
    main()
